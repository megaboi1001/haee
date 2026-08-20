"""Build the end-user Client VPN quick-start as a Word document.

Converts CLIENT-VPN-QUICKSTART.md (repo root) into
deliverables/Client-VPN-Quick-Start.docx so the client instruction ships as a
single Word file alongside the other handover documents.

Usage:  python3 deliverables/_build_client_quickstart.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
MD = ROOT / "CLIENT-VPN-QUICKSTART.md"
OUT = Path(__file__).resolve().parent / "Client-VPN-Quick-Start.docx"

NAVY = RGBColor(0x1F, 0x38, 0x64)
GRAY = RGBColor(0x59, 0x59, 0x59)


def add_runs(paragraph, text: str) -> None:
    """Add runs to a paragraph, honoring **bold**, *italic* and `code`."""
    token_re = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)")
    for tok in token_re.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            run = paragraph.add_run(tok[1:-1])
            run.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(tok)


def build() -> None:
    lines = MD.read_text(encoding="utf-8").splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        line = raw.strip()

        # horizontal rule -> spacer
        if line == "---":
            i += 1
            continue

        # heading
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = min(len(m.group(1)) + 1, 4)  # docx h1..h4
            text = m.group(2)
            h = doc.add_heading(level=level)
            add_runs(h, text)
            for run in h.runs:
                run.font.color.rgb = NAVY
            i += 1
            continue

        # blockquote
        if line.startswith(">"):
            txt = line.lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            add_runs(p, txt)
            for run in p.runs:
                run.font.color.rgb = GRAY
            i += 1
            continue

        # table: detect a header row followed by a separator row
        if line.startswith("|") and i + 1 < len(lines):
            sep = lines[i + 1].strip()
            if re.match(r"^\|[\s:|-]+\|?$", sep):
                header_row = [c.strip() for c in line.strip("|").split("|")]
                rows = []
                j = i + 2
                while j < len(lines) and lines[j].strip().startswith("|"):
                    rows.append(
                        [c.strip() for c in lines[j].strip().strip("|").split("|")]
                    )
                    j += 1
                table = doc.add_table(rows=1, cols=len(header_row))
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                for k, cell_text in enumerate(header_row):
                    hdr[k].text = ""
                    p = hdr[k].paragraphs[0]
                    r = p.add_run(cell_text)
                    r.bold = True
                for row in rows:
                    cells = table.add_row().cells
                    for k, cell_text in enumerate(row):
                        if k < len(cells):
                            cells[k].text = ""
                            add_runs(cells[k].paragraphs[0], cell_text)
                doc.add_paragraph()
                i = j
                continue

        # ordered list
        m = re.match(r"^(\d+)\.\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(2))
            i += 1
            continue

        # unordered list
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line[2:])
            i += 1
            continue

        # regular paragraph (skip empty lines)
        if line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_runs(p, line)
        i += 1

    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
